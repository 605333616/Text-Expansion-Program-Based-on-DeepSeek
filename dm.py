import tkinter as tk
from tkinter import filedialog, BOTTOM,ttk
from docx import Document
from ollama import chat
from ollama import ChatResponse





def select_file(i=0):
    label["text"] = "正在运行请稍等"
    # 弹出文件选择对话框，并获取选中的文件路径
    selected_file_path = filedialog.askopenfilename()
    doc = Document(selected_file_path)


    # 初始化一个空列表来存储文字内容
    text_content = []
    # 遍历每个段落
    for paragraph in doc.paragraphs:
        # 获取段落文本
        content = paragraph.text
        # 将回车符号替换为空格
        content = content.replace('\r\n', ' ')
        text_content.append(content)

    # 将所有段落内容合并成一个字符串
    #full_text = '\n'.join(text_content)

    # 输出到控制台
    #print(text_content[1])
    modelVersion=com.get()
    for element in text_content:
        response: ChatResponse = chat(model=modelVersion, messages=[
            {
                'role': 'system',
                'content': '只填充细节，不扩充情节,要有画面感，尽量简洁,限制在一个自然段以内：',
            },
            {
                "role": "user",
                "content": element
            }
        ])
        unThink=response.message.content
        unThink=unThink.split('</think>',1)
        text_content[i]=unThink[1]
        i=i+1

    full_text = '\n'.join(text_content)
    # 实例化一个Document对象，相当于打开word软件，新建一个空白文件
    doc = Document()
    # word文件尾部增加一个段落，并写入内容
    paragraph = doc.add_paragraph(full_text)
    # 保存word文件到当前文件夹
    address=selected_file_path.split('.',1)
    doc.save(address[0]+'-已修改版本'+'.'+address[1])
    label["text"] = "生成完成"


# 创建Tkinter窗口
root = tk.Tk()
root.title("文本扩充器")
root.geometry("500x400")



xVariable = tk.StringVar()  # #创建变量，便于取值

com = ttk.Combobox(root, textvariable=xVariable)  # #创建下拉菜单

com.pack()  # #将下拉菜单绑定到窗体
com["value"] = ("deepseek-r1:1.5b", "deepseek-r1:7b", "deepseek-r1:8b",'deepseek-r1:14b','deepseek-r1:32b')  # #给下拉菜单设定值
com.current(3)  #





# 创建并配置按钮，点击时调用select_file函数
button = tk.Button(root, text="选择文件", command=select_file)
button.pack()

label = tk.Label(root, text="选择文件开始扩充文本",font=('Calibri 15 bold'))
label.pack(pady=80)
statement=tk.Label(root, text="本项目完全免费，\n用户可以自由下载、使用以及分发本项目的代码和相关资源。\n你无需支付任何费用即可享受本项目的功能和服务。\n开发者不对因使用本项目而产生的任何直接或间接损失承担责任。\n用户在使用本项目时应自行承担风险，并对可能的后果负责。",
                   font=('Arial', 10))
statement.pack(side=BOTTOM)


# 运行Tkinter事件循环
root.mainloop()

